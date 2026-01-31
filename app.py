from flask import Flask, render_template, request, jsonify
from transformers import AutoTokenizer, AutoModel, pipeline
from PyPDF2 import PdfReader
from docx import Document
import torch
import torch.nn.functional as F
import os
import re
from pymongo import MongoClient
from datetime import datetime, timedelta
from bson import ObjectId
import jwt
import bcrypt
from functools import wraps
import string

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'your-secret-key-change-in-production-2024')
JWT_EXPIRATION_HOURS = 24 * 7  # 7 days

# ========================================
# MongoDB Database Setup for Chat History
# ========================================
MONGO_URI = 'mongodb://localhost:27017/'
DB_NAME = 'academic_qa_db'

# Cached MongoDB client (connection pooling)
_mongo_client = None

def get_db():
    """Get MongoDB database connection with connection pooling."""
    global _mongo_client
    if _mongo_client is None:
        _mongo_client = MongoClient(MONGO_URI, maxPoolSize=10)
    return _mongo_client[DB_NAME]

def init_db():
    """Initialize MongoDB collections with indexes."""
    try:
        db = get_db()
        
        # Create indexes for better performance
        db.chat_sessions.create_index('session_id', unique=True)
        db.chat_sessions.create_index('user_id')
        db.chat_messages.create_index('session_id')
        db.chat_messages.create_index('created_at')
        
        # User collection indexes
        db.users.create_index('email', unique=True)
        db.users.create_index('username', unique=True)
        
        # User documents collection
        db.user_documents.create_index('user_id')
        
        # Token blacklist for logout (with TTL to auto-expire old tokens)
        db.token_blacklist.create_index('token', unique=True)
        db.token_blacklist.create_index('expires_at', expireAfterSeconds=0)
        
        print("MongoDB connection established successfully!")
    except Exception as e:
        print(f"Warning: MongoDB connection failed - {e}")
        print("Chat history will not be available until MongoDB is running.")

# Initialize database on startup
init_db()

# ========================================
# Authentication Helper Functions
# ========================================

def hash_password(password):
    """Hash a password using bcrypt."""
    return bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')

def verify_password(password, hashed):
    """Verify a password against its hash."""
    return bcrypt.checkpw(password.encode('utf-8'), hashed.encode('utf-8'))

def generate_token(user_id, email):
    """Generate a JWT token for a user."""
    payload = {
        'user_id': str(user_id),
        'email': email,
        'exp': datetime.utcnow() + timedelta(hours=JWT_EXPIRATION_HOURS),
        'iat': datetime.utcnow()
    }
    return jwt.encode(payload, app.config['SECRET_KEY'], algorithm='HS256')

def decode_token(token):
    """Decode and verify a JWT token."""
    try:
        # First check if token is blacklisted
        db = get_db()
        if db.token_blacklist.find_one({'token': token}):
            return None
        
        payload = jwt.decode(token, app.config['SECRET_KEY'], algorithms=['HS256'])
        return payload
    except jwt.ExpiredSignatureError:
        return None
    except jwt.InvalidTokenError:
        return None

def blacklist_token(token):
    """Add a token to the blacklist."""
    try:
        # Decode to get expiration time
        payload = jwt.decode(token, app.config['SECRET_KEY'], algorithms=['HS256'], options={'verify_exp': False})
        exp_timestamp = payload.get('exp', 0)
        expires_at = datetime.utcfromtimestamp(exp_timestamp)
        
        db = get_db()
        db.token_blacklist.update_one(
            {'token': token},
            {'$set': {'token': token, 'expires_at': expires_at, 'blacklisted_at': datetime.utcnow()}},
            upsert=True
        )
        return True
    except Exception as e:
        print(f"Failed to blacklist token: {e}")
        return False

def get_current_user():
    """Get the current user from the request token."""
    auth_header = request.headers.get('Authorization')
    if not auth_header or not auth_header.startswith('Bearer '):
        return None
    
    token = auth_header.split(' ')[1]
    payload = decode_token(token)
    if not payload:
        return None
    
    db = get_db()
    user = db.users.find_one({'_id': ObjectId(payload['user_id'])})
    return user

def token_required(f):
    """Decorator to require authentication."""
    @wraps(f)
    def decorated(*args, **kwargs):
        user = get_current_user()
        if not user:
            return jsonify({'error': 'Authentication required'}), 401
        return f(user, *args, **kwargs)
    return decorated

def token_optional(f):
    """Decorator for optional authentication."""
    @wraps(f)
    def decorated(*args, **kwargs):
        user = get_current_user()
        return f(user, *args, **kwargs)
    return decorated

# Initialize the model for semantic embeddings
print("Loading embedding model... This may take a moment.")
model_name = 'sentence-transformers/all-MiniLM-L6-v2'
tokenizer = AutoTokenizer.from_pretrained(model_name)
model = AutoModel.from_pretrained(model_name)
model.eval()
print("Embedding model loaded successfully!")

# ========================================
# RoBERTa Extractive QA Model
# ========================================
print("Loading RoBERTa QA model for extractive question answering...")
qa_model = pipeline(
    "question-answering",
    model="deepset/roberta-base-squad2",
    tokenizer="deepset/roberta-base-squad2"
)
print("RoBERTa QA model loaded successfully!")

# QA Configuration
QA_CONFIG = {
    'min_confidence': 0.15,           # Minimum confidence threshold for answers
    'high_confidence': 0.5,           # High confidence threshold
    'max_answer_length': 200,         # Maximum answer length in characters
    'min_answer_length': 10,          # Minimum answer length in characters
    'chunk_size': 400,                # Optimal chunk size (350-450 tokens)
    'chunk_overlap': 50,              # Token overlap between chunks
    'top_k_chunks': 3,                # Number of top chunks to retrieve
    'context_window': 512,            # Max context window for QA model
}

# Global storage for document embeddings (trained on user's document)
document_store = {
    'chunks': [],           # Original text chunks
    'embeddings': None,     # Embeddings of chunks (the "trained" knowledge)
    'filename': None        # Current document name
}

def mean_pooling(model_output, attention_mask):
    """Mean pooling to get sentence embeddings."""
    token_embeddings = model_output[0]
    input_mask_expanded = attention_mask.unsqueeze(-1).expand(token_embeddings.size()).float()
    return torch.sum(token_embeddings * input_mask_expanded, 1) / torch.clamp(input_mask_expanded.sum(1), min=1e-9)

def get_embeddings(texts):
    """Get embeddings for a list of texts."""
    encoded_input = tokenizer(texts, padding=True, truncation=True, max_length=512, return_tensors='pt')
    
    with torch.no_grad():
        model_output = model(**encoded_input)
    
    embeddings = mean_pooling(model_output, encoded_input['attention_mask'])
    embeddings = F.normalize(embeddings, p=2, dim=1)
    return embeddings

def extract_text_from_pdf(file):
    """Extract text from a PDF file with better formatting preservation."""
    try:
        pdf_reader = PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n\n"
        return text.strip()
    except Exception as e:
        raise Exception(f"Failed to extract text from PDF: {str(e)}")

def extract_text_from_docx(file):
    """Extract text from a Word document."""
    try:
        doc = Document(file)
        text = ""
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n\n"
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text += " | ".join(row_text) + "\n"
        return text.strip()
    except Exception as e:
        raise Exception(f"Failed to extract text from Word document: {str(e)}")

def chunk_text(text, chunk_size=400, overlap=50):
    """
    Split text into overlapping chunks optimized for semantic search and QA.
    Uses token-based splitting (350-450 tokens) with section awareness.
    Respects section boundaries to avoid mixing content from different sections.
    """
    chunks = []
    
    # First, split by section headers (e.g., "2.1", "2.2", "Chapter 1", etc.)
    section_pattern = r'\n(?=\d+\.\d+\s|\d+\.\s|Chapter\s+\d+|Section\s+\d+|[A-Z][A-Z\s]+:|\n[A-Z][a-z]+[-\s]Oriented|\n\*\*)'
    sections = re.split(section_pattern, text, flags=re.IGNORECASE)
    
    for section in sections:
        section = section.strip()
        if not section:
            continue
        
        # Extract section header for context preservation
        section_header = ""
        header_match = re.match(r'^(\d+\.\d+\s+[^\n]+|\d+\.\s+[^\n]+|[A-Z][a-z]+[-\s]Oriented Phase[^\n]*)', section)
        if header_match:
            section_header = header_match.group(0).strip()
        
        # Tokenize section to count tokens
        section_tokens = tokenizer.encode(section, add_special_tokens=False)
        
        # If section fits within chunk_size, keep it as one chunk
        if len(section_tokens) <= chunk_size:
            if len(section) > 50:  # Minimum character threshold
                chunks.append(section)
            continue
        
        # For larger sections, split into optimal chunks with overlap
        # Split by sentences for cleaner boundaries
        sentences = re.split(r'(?<=[.!?])\s+', section)
        
        current_chunk_tokens = []
        current_chunk_text = []
        
        # Add section header to first chunk for context
        if section_header:
            header_tokens = tokenizer.encode(f"[{section_header}] ", add_special_tokens=False)
            current_chunk_tokens.extend(header_tokens)
            current_chunk_text.append(f"[{section_header}]")
        
        for sent in sentences:
            sent = sent.strip()
            if not sent:
                continue
            
            sent_tokens = tokenizer.encode(sent, add_special_tokens=False)
            
            # Check if adding this sentence exceeds chunk size
            if len(current_chunk_tokens) + len(sent_tokens) > chunk_size:
                # Save current chunk if it has content
                if current_chunk_text and len(' '.join(current_chunk_text)) > 50:
                    chunks.append(' '.join(current_chunk_text))
                
                # Start new chunk with overlap (last few sentences)
                overlap_tokens = 0
                overlap_start = len(current_chunk_text)
                
                # Find overlap point
                for i in range(len(current_chunk_text) - 1, -1, -1):
                    sent_len = len(tokenizer.encode(current_chunk_text[i], add_special_tokens=False))
                    if overlap_tokens + sent_len > overlap:
                        break
                    overlap_tokens += sent_len
                    overlap_start = i
                
                # New chunk starts with section header (for context) and overlap
                current_chunk_text = []
                current_chunk_tokens = []
                
                if section_header:
                    header_tokens = tokenizer.encode(f"[{section_header}] ", add_special_tokens=False)
                    current_chunk_tokens.extend(header_tokens)
                    current_chunk_text.append(f"[{section_header}]")
            
            current_chunk_tokens.extend(sent_tokens)
            current_chunk_text.append(sent)
        
        # Don't forget the last chunk
        if current_chunk_text and len(' '.join(current_chunk_text)) > 50:
            chunks.append(' '.join(current_chunk_text))
    
    # Filter out duplicates and very small chunks
    seen = set()
    unique_chunks = []
    for c in chunks:
        c_clean = c.strip()
        # Normalize for duplicate detection
        c_normalized = ' '.join(c_clean.lower().split())
        if len(c_clean) > 50 and c_normalized not in seen:
            seen.add(c_normalized)
            unique_chunks.append(c_clean)
    
    return unique_chunks

def train_on_document(text, filename):
    """
    'Train' the model on the document by creating embeddings.
    This is the core of semantic search - we encode all chunks once.
    """
    global document_store
    
    # Split text into meaningful chunks
    chunks = chunk_text(text)
    
    if not chunks:
        raise Exception("Could not extract meaningful content from the document.")
    
    # Create embeddings for all chunks (this is the "training" step)
    print(f"Training on document: {filename}")
    print(f"Creating embeddings for {len(chunks)} text chunks...")
    
    # Process in batches to avoid memory issues
    batch_size = 32
    all_embeddings = []
    
    for i in range(0, len(chunks), batch_size):
        batch = chunks[i:i + batch_size]
        batch_embeddings = get_embeddings(batch)
        all_embeddings.append(batch_embeddings)
        print(f"  Processed {min(i + batch_size, len(chunks))}/{len(chunks)} chunks...")
    
    embeddings = torch.cat(all_embeddings, dim=0)
    
    # Store the trained knowledge
    document_store['chunks'] = chunks
    document_store['embeddings'] = embeddings
    document_store['filename'] = filename
    
    print(f"Training complete! Model is now trained on '{filename}'")
    
    return len(chunks)

def find_relevant_chunks(question, top_k=3):
    """
    Find the most relevant chunks from the trained document.
    Uses cosine similarity between question and chunk embeddings.
    Also applies keyword boosting for more precise matching.
    """
    global document_store
    
    if document_store['embeddings'] is None:
        return []
    
    # Encode the question
    question_embedding = get_embeddings([question])
    
    # Calculate cosine similarity with all chunks using PyTorch
    similarities = F.cosine_similarity(question_embedding, document_store['embeddings'])
    
    # Extract key terms from question for boosting
    question_lower = question.lower()
    
    # Boost scores for chunks that contain specific section/topic keywords from the question
    boosted_similarities = similarities.clone()
    
    for idx, chunk in enumerate(document_store['chunks']):
        chunk_lower = chunk.lower()
        boost = 0.0
        
        # Check for section-specific keywords (e.g., "debugging-oriented", "demonstration-oriented")
        # Extract potential section identifiers from the question
        section_patterns = [
            r'(\d+\.\d+)',  # e.g., "2.1", "2.2"
            r'(debugging[- ]oriented)',
            r'(demonstration[- ]oriented)',
            r'(destruction[- ]oriented)',
            r'(prevention[- ]oriented)',
            r'(phase\s*\d+)',
            r'(chapter\s*\d+)',
        ]
        
        for pattern in section_patterns:
            matches = re.findall(pattern, question_lower)
            for match in matches:
                if match in chunk_lower:
                    boost += 0.15  # Significant boost for section match
        
        # Penalize chunks that mention OTHER sections when question is specific
        if 'debugging-oriented' in question_lower or 'debugging oriented' in question_lower or '2.1' in question_lower:
            if ('demonstration-oriented' in chunk_lower or 'demonstration oriented' in chunk_lower or 
                '2.2' in chunk_lower or 'destruction-oriented' in chunk_lower):
                boost -= 0.2  # Penalize wrong section
        
        if 'demonstration-oriented' in question_lower or 'demonstration oriented' in question_lower or '2.2' in question_lower:
            if ('debugging-oriented' in chunk_lower or 'debugging oriented' in chunk_lower or 
                '2.1' in chunk_lower):
                boost -= 0.2  # Penalize wrong section
        
        boosted_similarities[idx] += boost
    
    # Get top-k most similar chunks
    top_k = min(top_k, len(document_store['chunks']))
    top_scores, top_indices = torch.topk(boosted_similarities, top_k)
    
    results = []
    for score, idx in zip(top_scores, top_indices):
        # Use original similarity for the score display
        original_score = float(similarities[idx.item()].item())
        results.append({
            'chunk': document_store['chunks'][idx.item()],
            'score': original_score,
            'index': int(idx.item())
        })
    
    return results

def convert_to_question_format(query):
    """
    FIX 3: Convert keyword-style queries to proper question format.
    This significantly improves extractive QA performance.
    
    Bad (generative style):
        - "common software testing methodologies"
        - "explain agile testing"
        - "list types of testing"
    
    Good (extractive style):
        - "What are the common software testing methodologies?"
        - "What is agile testing?"
        - "What are the types of testing?"
    """
    query = query.strip()
    
    # If already ends with question mark, likely already a question
    if query.endswith('?'):
        return query
    
    # Check if already starts with a question word
    question_starters = ['what', 'who', 'where', 'when', 'why', 'how', 'which', 'is', 'are', 'do', 'does', 'can', 'could', 'would', 'should', 'will', 'did', 'has', 'have', 'was', 'were']
    query_lower = query.lower()
    
    if any(query_lower.startswith(starter + ' ') for starter in question_starters):
        return query + '?'
    
    # Convert imperative/keyword style to question style
    imperative_mappings = {
        'explain': 'What is',
        'describe': 'What is',
        'define': 'What is',
        'list': 'What are the',
        'give': 'What are',
        'tell': 'What is',
        'show': 'What is',
        'find': 'What is',
        'get': 'What is',
    }
    
    for imperative, question_form in imperative_mappings.items():
        if query_lower.startswith(imperative + ' '):
            remaining = query[len(imperative):].strip()
            return f"{question_form} {remaining}?"
    
    # For simple keyword queries, convert to "What is/are X?"
    # Check if query seems plural
    words = query.split()
    if len(words) > 0:
        # Simple heuristic: if ends with 's' or contains 'types', 'methods', etc.
        plural_indicators = ['types', 'methods', 'methodologies', 'techniques', 'steps', 'phases', 'stages', 'kinds', 'forms', 'categories', 'levels', 'principles', 'advantages', 'disadvantages', 'features', 'characteristics', 'examples']
        
        if any(indicator in query_lower for indicator in plural_indicators):
            return f"What are the {query}?"
        else:
            return f"What is {query}?"
    
    return query + '?'

# ========================================
# Extractive QA and Post-Processing
# ========================================

def extract_answer_with_roberta(question, context, max_length=None):
    """
    Use RoBERTa model to extract precise answer spans from context.
    Returns answer, score, and start/end positions.
    """
    if max_length is None:
        max_length = QA_CONFIG['max_answer_length']
    
    try:
        result = qa_model(
            question=question,
            context=context,
            max_answer_len=max_length,
            handle_impossible_answer=True  # For SQuAD 2.0 "no answer" detection
        )
        
        return {
            'answer': result['answer'],
            'score': result['score'],
            'start': result['start'],
            'end': result['end']
        }
    except Exception as e:
        print(f"RoBERTa extraction error: {e}")
        return None

def post_process_answer(answer, question, min_length=None, max_length=None):
    """
    Post-process extracted answer to improve quality:
    1. Trim extraneous sentences
    2. Remove incomplete sentences at boundaries
    3. Clean up formatting
    4. Ensure answer is relevant to question
    """
    if min_length is None:
        min_length = QA_CONFIG['min_answer_length']
    if max_length is None:
        max_length = QA_CONFIG['max_answer_length']
    
    if not answer or len(answer.strip()) < min_length:
        return None
    
    answer = answer.strip()
    
    # Remove leading/trailing incomplete sentences
    # Check for incomplete start (no capital letter or starts mid-sentence)
    if answer and not answer[0].isupper() and answer[0] not in string.digits:
        # Find first complete sentence
        first_period = answer.find('. ')
        if first_period > 0 and first_period < len(answer) - 10:
            answer = answer[first_period + 2:].strip()
    
    # Check for incomplete end (doesn't end with proper punctuation)
    if answer and answer[-1] not in '.!?':
        # Find last complete sentence
        last_period = answer.rfind('. ')
        if last_period > 0:
            answer = answer[:last_period + 1].strip()
    
    # Truncate if too long while preserving sentence boundaries
    if len(answer) > max_length:
        # Find the last sentence boundary before max_length
        truncated = answer[:max_length]
        last_period = truncated.rfind('. ')
        if last_period > max_length * 0.5:  # At least half the content
            answer = truncated[:last_period + 1].strip()
        else:
            answer = truncated.strip() + '...'
    
    # Clean up extra whitespace
    answer = ' '.join(answer.split())
    
    return answer if len(answer) >= min_length else None

def is_answer_relevant(answer, question, context):
    """
    Check if the extracted answer is relevant to the question.
    Uses simple heuristics to filter out noise.
    """
    if not answer:
        return False
    
    answer_lower = answer.lower()
    question_lower = question.lower()
    
    # Extract key terms from question
    question_words = set(re.findall(r'\b[a-z]{3,}\b', question_lower))
    stop_words = {'what', 'which', 'where', 'when', 'why', 'how', 'does', 'the', 'and', 'for', 'are', 'was', 'were', 'been', 'being', 'have', 'has', 'had', 'having', 'this', 'that', 'these', 'those', 'there'}
    question_keywords = question_words - stop_words
    
    # Check if answer contains at least one question keyword
    answer_words = set(re.findall(r'\b[a-z]{3,}\b', answer_lower))
    
    # Special case: if question is about a specific topic, answer should relate
    overlap = question_keywords & answer_words
    
    # Allow answers that are explanatory even without keyword overlap
    if len(answer) > 30 and answer in context:
        return True
    
    return len(overlap) > 0 or len(answer) > 50

def two_stage_qa_pipeline(question, top_chunks):
    """
    Two-stage QA pipeline:
    Stage 1: Semantic retrieval (already done - top_chunks provided)
    Stage 2: Extractive QA on each retrieved chunk using RoBERTa
    
    Returns the best answer with confidence and supporting context.
    """
    if not top_chunks:
        return None
    
    candidates = []
    
    for chunk_data in top_chunks:
        chunk = chunk_data['chunk']
        retrieval_score = chunk_data['score']
        
        # Extract answer using RoBERTa
        qa_result = extract_answer_with_roberta(question, chunk)
        
        if qa_result and qa_result['answer']:
            # Post-process the answer
            processed_answer = post_process_answer(
                qa_result['answer'], 
                question
            )
            
            if processed_answer and is_answer_relevant(processed_answer, question, chunk):
                # Combined score: QA confidence + retrieval score
                combined_score = (qa_result['score'] * 0.7) + (retrieval_score * 0.3)
                
                candidates.append({
                    'answer': processed_answer,
                    'raw_answer': qa_result['answer'],
                    'qa_score': qa_result['score'],
                    'retrieval_score': retrieval_score,
                    'combined_score': combined_score,
                    'context': chunk,
                    'chunk_index': chunk_data['index']
                })
    
    if not candidates:
        # Fallback: return best chunk content if no extractive answer found
        return {
            'answer': None,
            'fallback': True,
            'context': top_chunks[0]['chunk'] if top_chunks else None,
            'retrieval_score': top_chunks[0]['score'] if top_chunks else 0
        }
    
    # Sort by combined score and return best
    candidates.sort(key=lambda x: x['combined_score'], reverse=True)
    best = candidates[0]
    
    # Include alternative answers if confidence is similar
    alternatives = []
    for c in candidates[1:3]:
        if c['combined_score'] > best['combined_score'] * 0.8:
            alternatives.append(c)
    
    return {
        'answer': best['answer'],
        'qa_score': best['qa_score'],
        'retrieval_score': best['retrieval_score'],
        'combined_score': best['combined_score'],
        'context': best['context'],
        'chunk_index': best['chunk_index'],
        'alternatives': alternatives,
        'fallback': False
    }

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    """Handle file upload and train the model on the document."""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        filename = file.filename.lower()
        original_filename = file.filename
        
        if filename.endswith('.pdf'):
            text = extract_text_from_pdf(file)
        elif filename.endswith('.docx'):
            text = extract_text_from_docx(file)
        elif filename.endswith('.doc'):
            return jsonify({'error': 'Old .doc format is not supported. Please convert to .docx format.'}), 400
        elif filename.endswith('.txt'):
            text = file.read().decode('utf-8')
        else:
            return jsonify({'error': 'Unsupported file type. Please upload PDF, DOCX, or TXT file.'}), 400
        
        if not text.strip():
            return jsonify({'error': 'Could not extract any text from the file.'}), 400
        
        # TRAIN the model on this document
        num_chunks = train_on_document(text, original_filename)
        
        return jsonify({
            'success': True,
            'text': text,
            'filename': original_filename,
            'message': f'Successfully trained on "{original_filename}" ({num_chunks} knowledge chunks created)'
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# ========================================
# Authentication API Endpoints
# ========================================

@app.route('/api/auth/register', methods=['POST'])
def register():
    """Register a new user."""
    try:
        data = request.get_json()
        username = data.get('username', '').strip()
        email = data.get('email', '').strip().lower()
        password = data.get('password', '')
        
        # Validation
        if not username or len(username) < 3:
            return jsonify({'error': 'Username must be at least 3 characters'}), 400
        if not email or '@' not in email:
            return jsonify({'error': 'Please provide a valid email'}), 400
        if not password or len(password) < 6:
            return jsonify({'error': 'Password must be at least 6 characters'}), 400
        
        db = get_db()
        
        # Check if user already exists
        if db.users.find_one({'email': email}):
            return jsonify({'error': 'Email already registered'}), 400
        if db.users.find_one({'username': username}):
            return jsonify({'error': 'Username already taken'}), 400
        
        # Create user
        now = datetime.utcnow()
        user = {
            'username': username,
            'email': email,
            'password': hash_password(password),
            'created_at': now,
            'updated_at': now,
            'preferences': {}
        }
        
        result = db.users.insert_one(user)
        user_id = result.inserted_id
        
        # Generate token
        token = generate_token(user_id, email)
        
        return jsonify({
            'success': True,
            'token': token,
            'user': {
                'id': str(user_id),
                'username': username,
                'email': email
            }
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/auth/login', methods=['POST'])
def login():
    """Login a user."""
    try:
        data = request.get_json()
        email = data.get('email', '').strip().lower()
        password = data.get('password', '')
        
        if not email or not password:
            return jsonify({'error': 'Email and password required'}), 400
        
        db = get_db()
        user = db.users.find_one({'email': email})
        
        if not user or not verify_password(password, user['password']):
            return jsonify({'error': 'Invalid email or password'}), 401
        
        # Generate token
        token = generate_token(user['_id'], email)
        
        return jsonify({
            'success': True,
            'token': token,
            'user': {
                'id': str(user['_id']),
                'username': user['username'],
                'email': user['email']
            }
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/auth/me', methods=['GET'])
@token_required
def get_me(user):
    """Get current user info."""
    return jsonify({
        'user': {
            'id': str(user['_id']),
            'username': user['username'],
            'email': user['email']
        }
    })

@app.route('/api/auth/logout', methods=['POST'])
def logout():
    """Logout - invalidate the token on server side."""
    auth_header = request.headers.get('Authorization')
    if auth_header and auth_header.startswith('Bearer '):
        token = auth_header.split(' ')[1]
        blacklist_token(token)
    
    return jsonify({'success': True, 'message': 'Logged out successfully'})

@app.route('/ask', methods=['POST'])
def ask():
    """
    Two-stage Question Answering Pipeline:
    Stage 1: Semantic retrieval - find top-k relevant chunks using embeddings
    Stage 2: Extractive QA - use RoBERTa to extract precise answer spans
    """
    try:
        data = request.get_json()
        question = data.get('question', '').strip()
        context = data.get('context', '').strip()
        
        if not question:
            return jsonify({'error': 'Please provide a question.'}), 400
        
        # Convert keyword-style queries to proper question format
        # This significantly improves extractive QA performance
        original_question = question
        question = convert_to_question_format(question)
        
        # Check if we have a trained document
        if document_store['embeddings'] is None:
            if context:
                train_on_document(context, "pasted_text")
            else:
                return jsonify({'error': 'Please upload a document first to train the model.'}), 400
        
        # ========================================
        # STAGE 1: Semantic Retrieval
        # ========================================
        relevant_chunks = find_relevant_chunks(question, top_k=QA_CONFIG['top_k_chunks'])
        
        if not relevant_chunks:
            return jsonify({
                'answer': 'Could not find relevant information in the document.',
                'score': 0,
                'is_fallback': True
            })
        
        # Check minimum retrieval threshold
        best_retrieval_score = relevant_chunks[0]['score']
        
        if best_retrieval_score < QA_CONFIG['min_confidence']:
            return jsonify({
                'answer': 'I couldn\'t find a confident answer to your question in the document. Please try rephrasing your question to be more specific (e.g., "What is X?" or "How does Y work?").',
                'score': best_retrieval_score,
                'is_fallback': True,
                'fallback_reason': 'low_retrieval_confidence'
            })
        
        # ========================================
        # STAGE 2: Extractive QA with RoBERTa
        # ========================================
        qa_result = two_stage_qa_pipeline(question, relevant_chunks)
        
        # Prepare source chunks for context display
        source_chunks = []
        for chunk in relevant_chunks:
            if chunk['score'] >= QA_CONFIG['min_confidence']:
                source_chunks.append({
                    'text': chunk['chunk'],
                    'score': round(chunk['score'] * 100, 1),
                    'index': chunk['index']
                })
        
        # Handle fallback case (no extractive answer found)
        if qa_result.get('fallback', False) or not qa_result.get('answer'):
            # Check if retrieval score is high enough to show context
            if best_retrieval_score >= QA_CONFIG['high_confidence']:
                # Return relevant context with explanation
                context_answer = relevant_chunks[0]['chunk']
                # Trim to reasonable length
                if len(context_answer) > 500:
                    context_answer = context_answer[:500] + '...'
                
                return jsonify({
                    'answer': context_answer,
                    'short_answer': 'The answer can be found in the context below.',
                    'score': best_retrieval_score,
                    'source': document_store['filename'],
                    'chunks_found': len(source_chunks),
                    'source_chunks': source_chunks,
                    'is_extractive': False,
                    'show_context': True
                })
            else:
                return jsonify({
                    'answer': 'I found some related information but couldn\'t extract a precise answer. Please try asking in a different way or be more specific.',
                    'score': best_retrieval_score,
                    'is_fallback': True,
                    'source_chunks': source_chunks,
                    'fallback_reason': 'no_extractive_answer'
                })
        
        # ========================================
        # HIGH CONFIDENCE EXTRACTIVE ANSWER
        # ========================================
        answer = qa_result['answer']
        qa_score = qa_result['qa_score']
        combined_score = qa_result['combined_score']
        
        # Apply confidence threshold
        if qa_score < QA_CONFIG['min_confidence']:
            # Low confidence - provide answer with context
            return jsonify({
                'answer': answer,
                'short_answer': answer,
                'score': combined_score,
                'qa_score': qa_score,
                'retrieval_score': qa_result['retrieval_score'],
                'source': document_store['filename'],
                'chunks_found': len(source_chunks),
                'source_chunks': source_chunks,
                'is_extractive': True,
                'confidence_level': 'low',
                'show_context': True  # Show context for low confidence
            })
        
        # Determine confidence level for UI
        confidence_level = 'high' if qa_score >= QA_CONFIG['high_confidence'] else 'medium'
        
        # Build response
        response = {
            'answer': answer,
            'short_answer': answer,
            'score': combined_score,
            'qa_score': qa_score,
            'retrieval_score': qa_result['retrieval_score'],
            'source': document_store['filename'],
            'chunks_found': len(source_chunks),
            'source_chunks': source_chunks,
            'is_extractive': True,
            'confidence_level': confidence_level,
            'show_context': confidence_level != 'high'  # Auto-show context for non-high confidence
        }
        
        # Add alternatives if available
        if qa_result.get('alternatives'):
            response['alternatives'] = [
                {
                    'answer': alt['answer'],
                    'score': alt['combined_score']
                }
                for alt in qa_result['alternatives']
            ]
        
        return jsonify(response)
        
    except Exception as e:
        import traceback
        print(f"Error in /ask: {e}")
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/status', methods=['GET'])
def status():
    """Get the current training status."""
    return jsonify({
        'trained': document_store['embeddings'] is not None,
        'filename': document_store['filename'],
        'num_chunks': len(document_store['chunks']) if document_store['chunks'] else 0
    })

# ========================================
# Chat History API Endpoints
# ========================================

@app.route('/api/sessions', methods=['GET'])
@token_optional
def get_sessions(user):
    """Get all chat sessions for the current user."""
    try:
        db = get_db()
        
        # Filter by user if authenticated
        query = {}
        if user:
            query['user_id'] = str(user['_id'])
        else:
            query['user_id'] = {'$exists': False}
        
        sessions = list(db.chat_sessions.find(query).sort('updated_at', -1).limit(50))
        
        result = []
        for session in sessions:
            result.append({
                'session_id': session['session_id'],
                'title': session['title'],
                'document_name': session.get('document_name'),
                'created_at': session['created_at'].isoformat() if session.get('created_at') else None,
                'updated_at': session['updated_at'].isoformat() if session.get('updated_at') else None
            })
        
        return jsonify({'sessions': result})
    except Exception as e:
        return jsonify({'sessions': [], 'error': str(e)})

@app.route('/api/sessions', methods=['POST'])
@token_optional
def create_session(user):
    """Create a new chat session."""
    try:
        data = request.get_json()
        session_id = data.get('session_id')
        title = data.get('title', 'New Chat')
        document_name = data.get('document_name')
        
        db = get_db()
        now = datetime.utcnow()
        
        session_data = {
            'session_id': session_id,
            'title': title,
            'document_name': document_name,
            'created_at': now,
            'updated_at': now
        }
        
        # Associate with user if authenticated
        if user:
            session_data['user_id'] = str(user['_id'])
        
        db.chat_sessions.insert_one(session_data)
        
        return jsonify({'success': True, 'session_id': session_id})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/sessions/<session_id>', methods=['PUT'])
def update_session(session_id):
    """Update a chat session title."""
    try:
        data = request.get_json()
        title = data.get('title')
        
        db = get_db()
        db.chat_sessions.update_one(
            {'session_id': session_id},
            {'$set': {'title': title, 'updated_at': datetime.utcnow()}}
        )
        
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/sessions/<session_id>', methods=['DELETE'])
def delete_session(session_id):
    """Delete a chat session and its messages."""
    try:
        db = get_db()
        db.chat_messages.delete_many({'session_id': session_id})
        db.chat_sessions.delete_one({'session_id': session_id})
        
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/sessions/<session_id>/messages', methods=['GET'])
def get_messages(session_id):
    """Get all messages for a session."""
    try:
        db = get_db()
        messages = list(db.chat_messages.find(
            {'session_id': session_id}
        ).sort('created_at', 1))
        
        result = []
        for msg in messages:
            result.append({
                'type': msg['message_type'],
                'content': msg['content'],
                'confidence': msg.get('confidence'),
                'created_at': msg['created_at'].isoformat() if msg.get('created_at') else None
            })
        
        return jsonify({'messages': result})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/sessions/<session_id>/messages', methods=['POST'])
def add_message(session_id):
    """Add a message to a session."""
    try:
        data = request.get_json()
        message_type = data.get('type')  # 'user' or 'bot'
        content = data.get('content')
        confidence = data.get('confidence')
        
        db = get_db()
        now = datetime.utcnow()
        
        # Insert message
        db.chat_messages.insert_one({
            'session_id': session_id,
            'message_type': message_type,
            'content': content,
            'confidence': confidence,
            'created_at': now
        })
        
        # Update session timestamp
        db.chat_sessions.update_one(
            {'session_id': session_id},
            {'$set': {'updated_at': now}}
        )
        
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/sessions/clear', methods=['DELETE'])
def clear_all_sessions():
    """Clear all chat history."""
    try:
        db = get_db()
        db.chat_messages.delete_many({})
        db.chat_sessions.delete_many({})
        
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(debug=False, host='0.0.0.0', port=5000)
